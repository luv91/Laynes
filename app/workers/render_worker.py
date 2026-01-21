"""
Render Worker

Converts raw documents to canonical line-numbered text for evidence tracking.

Output format:
L0001: First line of document
L0002: Second line
...

This enables evidence citations like "Lines L0047-L0052"
"""

import io
import logging
import re
import xml.etree.ElementTree as ET
from datetime import datetime
from typing import Optional

from bs4 import BeautifulSoup

from app.web.db import db
from app.models.document_store import OfficialDocument
from app.models.ingest_job import IngestJob

logger = logging.getLogger(__name__)


class RenderWorker:
    """
    Converts raw documents to canonical line-numbered text.

    Supports:
    - XML (Federal Register with <GPOTABLE>)
    - HTML (CBP CSMS bulletins)
    - PDF (using pdfplumber)
    - DOCX (using python-docx)
    """

    def process(self, doc: OfficialDocument, job: IngestJob = None) -> bool:
        """
        Render document to canonical text.

        Args:
            doc: OfficialDocument with raw_bytes
            job: Optional IngestJob for status tracking

        Returns:
            True if successful
        """
        if job:
            job.status = "rendering"
            db.session.commit()

        try:
            content_type = doc.content_type or ""
            # Use doc.content property (reads from storage_uri or legacy raw_bytes)
            raw_content = doc.content

            if "xml" in content_type:
                canonical_text = self._render_xml(raw_content)
            elif "html" in content_type:
                canonical_text = self._render_html(raw_content)
            elif "pdf" in content_type:
                canonical_text = self._render_pdf(raw_content)
            elif "wordprocessingml" in content_type or "docx" in content_type.lower():
                canonical_text = self._render_docx(raw_content)
            else:
                # Default to HTML parsing
                canonical_text = self._render_html(raw_content)

            doc.canonical_text = canonical_text
            doc.status = "rendered"
            doc.rendered_at = datetime.utcnow()

            if job:
                job.status = "rendered"

            db.session.commit()

            logger.info(f"Rendered {doc.external_id}: {len(canonical_text)} chars, "
                       f"{canonical_text.count(chr(10))} lines")
            return True

        except Exception as e:
            logger.error(f"Render failed for {doc.external_id}: {e}")
            if job:
                job.mark_failed(f"Render error: {e}")
            db.session.commit()
            return False

    def _render_xml(self, raw_bytes: bytes) -> str:
        """
        Parse Federal Register XML.

        Extracts text from:
        - <P>: Paragraphs
        - <FP>: Formatted paragraphs
        - <GPOTABLE>: Structured tables
        - <HD>: Headings
        - <E>: Emphasis
        """
        try:
            root = ET.fromstring(raw_bytes)
        except ET.ParseError:
            # Try with encoding declaration stripped
            text = raw_bytes.decode('utf-8', errors='ignore')
            text = re.sub(r'<\?xml[^>]+\?>', '', text)
            root = ET.fromstring(text.encode('utf-8'))

        lines = []
        line_num = 1

        def extract_text(elem, depth=0):
            nonlocal line_num, lines

            # Get element tag
            tag = elem.tag if elem.tag else ""

            # Handle block elements (add line breaks)
            block_tags = {'P', 'FP', 'HD', 'ROW', 'GPOTABLE', 'SIG', 'DATE', 'AGENCY'}

            if tag in block_tags and elem.text:
                text = elem.text.strip()
                if text:
                    for text_line in text.split('\n'):
                        text_line = text_line.strip()
                        if text_line:
                            lines.append(f"L{line_num:04d}: {text_line}")
                            line_num += 1

            # Handle table rows
            if tag == 'ROW':
                entries = [e.text or "" for e in elem.findall('.//ENT')]
                if any(entries):
                    row_text = " | ".join(e.strip() for e in entries if e.strip())
                    lines.append(f"L{line_num:04d}: {row_text}")
                    line_num += 1

            # Handle headings specially
            if tag == 'HD':
                source = elem.get('SOURCE', '')
                text = elem.text or ""
                if text.strip():
                    lines.append(f"L{line_num:04d}: === {text.strip()} ===")
                    line_num += 1

            # Recurse into children
            for child in elem:
                extract_text(child, depth + 1)

            # Handle tail text
            if elem.tail and elem.tail.strip():
                for text_line in elem.tail.strip().split('\n'):
                    text_line = text_line.strip()
                    if text_line:
                        lines.append(f"L{line_num:04d}: {text_line}")
                        line_num += 1

        extract_text(root)
        return '\n'.join(lines)

    def _render_html(self, raw_bytes: bytes) -> str:
        """
        Parse HTML content (CBP bulletins, etc).
        """
        try:
            text = raw_bytes.decode('utf-8', errors='ignore')
        except:
            text = raw_bytes.decode('latin-1', errors='ignore')

        soup = BeautifulSoup(text, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()

        # Get text
        content = soup.get_text(separator='\n', strip=True)

        # Line number
        lines = []
        line_num = 1

        for line in content.split('\n'):
            line = line.strip()
            if line:
                lines.append(f"L{line_num:04d}: {line}")
                line_num += 1

        return '\n'.join(lines)

    def _render_pdf(self, raw_bytes: bytes) -> str:
        """
        Extract text from PDF using pdfplumber.
        """
        try:
            import pdfplumber
        except ImportError:
            logger.error("pdfplumber not installed")
            return ""

        lines = []
        line_num = 1

        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Add page marker
                    lines.append(f"L{line_num:04d}: === PAGE {page_num} ===")
                    line_num += 1

                    text = page.extract_text() or ""
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            lines.append(f"L{line_num:04d}: {line}")
                            line_num += 1

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")

        return '\n'.join(lines)

    def _render_docx(self, raw_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx.
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return ""

        lines = []
        line_num = 1

        try:
            doc = Document(io.BytesIO(raw_bytes))

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            lines.append(f"L{line_num:04d}: {line}")
                            line_num += 1

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        row_text = " | ".join(c for c in cells if c)
                        lines.append(f"L{line_num:04d}: {row_text}")
                        line_num += 1

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")

        return '\n'.join(lines)
